#!/usr/bin/env python3
"""report_to_docx.py: render REPORT.md (+ inlined build briefs) to REPORT.docx.

Handles the constructs this report actually uses: #/##/### headings, hard-wrapped
paragraphs, "- " bullets, "N. " numbered items, and pipe tables. Pure python-docx.
"""
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.join(HERE, "..")


def demote(md, levels=2):
    out = []
    for line in md.splitlines():
        if re.match(r"^#{1,4} ", line):
            out.append("#" * levels + line)
        else:
            out.append(line)
    return "\n".join(out)


def combined_markdown():
    t = open(os.path.join(ROOT, "REPORT.md"), encoding="utf-8").read()
    briefs = []
    for i, fn in enumerate([
        "brief-1-oath-ecb-defense-packet.md",
        "brief-2-snf-pbj-guardrail.md",
        "brief-3-fisp-deadline-radar.md",
    ], 1):
        b = open(os.path.join(ROOT, "briefs", fn), encoding="utf-8").read()
        briefs.append(demote(b, 2))
    marker = "Three ready-to-run Fable build briefs (full text in briefs/):"
    inline = ("The three ready-to-run Fable build briefs follow in full "
              "(also at briefs/ in the repository).\n\n" + "\n\n".join(briefs))
    if marker in t:
        # replace the pointer list (marker + 3 lines) with the full briefs
        t = re.sub(re.escape(marker) + r"\n(?:\d\. .*\n?){3}", inline + "\n", t)
    else:
        t += "\n\n" + inline
    return t


def blocks(md):
    """Yield (kind, payload) blocks: heading/table/bullets/numbers/para."""
    lines = md.splitlines()
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        m = re.match(r"^(#{1,6}) (.*)$", line)
        if m:
            yield ("heading", (len(m.group(1)), m.group(2).strip()))
            i += 1
            continue
        if line.lstrip().startswith("|"):
            rows = []
            while i < n and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            yield ("table", rows)
            continue
        if line.startswith("- "):
            items = []
            while i < n and (lines[i].startswith("- ") or (items and lines[i].startswith("  ") and lines[i].strip())):
                if lines[i].startswith("- "):
                    items.append(lines[i][2:].strip())
                else:
                    items[-1] += " " + lines[i].strip()
                i += 1
            yield ("bullets", items)
            continue
        if re.match(r"^\d{1,2}\. ", line):
            items = []
            while i < n and lines[i].strip():
                m2 = re.match(r"^(\d{1,2})\. (.*)$", lines[i])
                if m2:
                    items.append(m2.group(2).strip())
                elif items:
                    items[-1] += " " + lines[i].strip()
                i += 1
            yield ("numbers", items)
            continue
        para = []
        while i < n and lines[i].strip() and not re.match(r"^(#{1,6} |\- |\d{1,2}\. )", lines[i]) and not lines[i].lstrip().startswith("|"):
            para.append(lines[i].strip())
            i += 1
        yield ("para", " ".join(para))


def main():
    md = combined_markdown()
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    title = doc.add_heading("Vertical AI Micro-SaaS Opportunity Scanner", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x36, 0x50)
    sub = doc.add_paragraph("NSigma research run, 2026-07-04. Evidence-backed: every source "
                            "carries a live verification receipt from the run session.")
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    first_h1_seen = False
    for kind, payload in blocks(md):
        if kind == "heading":
            lvl, text = payload
            if lvl == 1 and not first_h1_seen:
                first_h1_seen = True
                continue  # markdown H1 duplicated by the docx title
            doc.add_heading(text, level=min(lvl, 4))
        elif kind == "para":
            if payload.startswith("NSigma run 2026-07-04. Branch"):
                continue  # covered by subtitle
            doc.add_paragraph(payload)
        elif kind == "bullets":
            for it in payload:
                doc.add_paragraph(it, style="List Bullet")
        elif kind == "numbers":
            for it in payload:
                doc.add_paragraph(it, style="List Number")
        elif kind == "table":
            rows = payload
            if not rows:
                continue
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            tbl.style = "Light Grid Accent 1"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                for ci in range(ncols):
                    cell = tbl.cell(ri, ci)
                    cell.text = row[ci] if ci < len(row) else ""
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(7.5)
                            if ri == 0:
                                r.font.bold = True
            doc.add_paragraph("")

    out = os.path.join(ROOT, "REPORT.docx")
    doc.save(out)
    print("wrote", out, os.path.getsize(out), "bytes")


if __name__ == "__main__":
    main()
